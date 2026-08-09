# -*- coding: utf-8 -*-
"""补充Q1和Q2的代码片段"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ===== Q1: Add PAGCM core code =====
print('Updating Q1...')
q1_path = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出\第一问_PAGCM完整报告_含图表.docx'
q1 = Document(q1_path)

for i, p in enumerate(q1.paragraphs):
    if '模块三：模型检验' in p.text:
        target = p._element
        code_lines = [
            '# ==== PAGCM核心类关键方法 ====',
            'class PAGCM:',
            '    def __init__(self, points_3d, r0=250, alpha=0.5, L=10000):',
            '        self.N = len(points_3d); self.r0 = r0',
            '        shift = L/2.0',
            '        self.pts = [(x+shift,y+shift,z+shift) for x,y,z in points_3d]',
            '        self.r_eff = [r0] * self.N',
            '        self.adj_edges = []',
            '',
            '# 环面距离(周期感知度量)',
            '@staticmethod',
            'def torus_dist(pi, pj, Lval):',
            '    d2 = 0.0',
            '    for dim in range(3):',
            '        diff = abs(pi[dim] - pj[dim])',
            '        d2 += min(diff, Lval - diff) ** 2',
            '    return math.sqrt(d2)',
            '',
            '# 密度感知自适应等效半径(PAGCM核心创新)',
            'def compute_adaptive_radius(self):',
            '    rho_global = self.N / (self.L ** 3)',
            '    for i in range(self.N):',
            '        pi = self.pts[i]; count = 0',
            '        for j in range(self.N):',
            '            if i == j: continue',
            '            if self.torus_dist(pi,self.pts[j],self.L) <= self.r_search:',
            '                count += 1',
            '        rho_local = count/(4.0/3.0*math.pi*self.r_search**3)',
            '        ratio = rho_local / max(rho_global, 1e-30)',
            '        re = self.r0 * (1.0 + self.alpha * (ratio - 1.0))',
            '        self.r_eff[i] = max(0.5*self.r0, min(3.0*self.r0, re))',
            '',
            '# 并查集聚类(路径压缩+按秩合并)',
            'def union_find_cluster(self):',
            '    parent, rank = list(range(self.N)), [0]*self.N',
            '    def find(x):',
            '        while parent[x] != x:',
            '            parent[x] = parent[parent[x]]; x = parent[x]',
            '        return x',
            '    def union(x, y):',
            '        rx, ry = find(x), find(y)',
            '        if rx == ry: return',
            '        if rank[rx] < rank[ry]: parent[rx] = ry',
            '        elif rank[rx] > rank[ry]: parent[ry] = rx',
            '        else: parent[ry] = rx; rank[rx] += 1',
            '    for i, j, kvec, d in self.adj_edges: union(i, j)',
            '',
            '# 方向性连通判定',
            'def check_connectivity(self, direction="X"):',
            '    axis = {"X":0, "Y":1, "Z":2}[direction]',
            '    lo = {i for i in range(self.N)',
            '          if self.pts[i][axis]-self.r_eff[i]<=0}',
            '    hi = {i for i in range(self.N)',
            '          if self.pts[i][axis]+self.r_eff[i]>=self.L}',
            '    for i in lo:',
            '        for j in hi:',
            '            if self.components[i]==self.components[j]: return True',
            '    return False',
        ]
        cap = q1.add_paragraph()
        r = cap.add_run('代码片段 Q1-1 PAGCM核心类关键方法 (完整代码见q1_solve.py, 约480行)')
        r.font.size = Pt(9); r.font.bold = True; r.font.italic = True
        target.addprevious(cap._element)
        for line in code_lines:
            cp = q1.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cp.add_run(line); r.font.size = Pt(8); r.font.name = 'Consolas'
            cp.space_after = Pt(2); cp.space_before = Pt(0)
            target.addprevious(cp._element)
        # Add blank line after code
        sp = q1.add_paragraph(); target.addprevious(sp._element)
        break

q1_v2 = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出\第一问_PAGCM完整报告_含图表_v2.docx'
q1.save(q1_v2)
print('[OK] Q1 v2: %.0f KB' % (os.path.getsize(q1_v2)/1024))

# ===== Q2: Add MESA algorithm code =====
print('Updating Q2...')
q2_path = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出\第二问_MESA-PAGCM完整报告_含图表_v2.docx'
q2 = Document(q2_path)

for i, p in enumerate(q2.paragraphs):
    if '模块三：模型检验' in p.text:
        target = p._element
        code_lines = [
            '# ==== MESA-PAGCM 模拟退火优化核心代码 ====',
            'def mesa_pagcm_optimize(P_target=0.95, T0=50, gamma=0.95):',
            '    # 1. MaxEnt初始化 (迁移1: 信息论->材料)',
            '    N = (N_min + N_max) // 2',
            '    particles = poisson_disk_sampling(N, L, beta=0.70)',
            '',
            '    # 2. 初始PAGCM评估',
            '    model = PAGCM(particles, r0=250, alpha=0.5, L=10000)',
            '    model.solve_quiet()',
            '    Pc = sum(model.connectivity.values())/3',
            '    f_curr = N/N_max + lambda_p*max(0, P_target-Pc)',
            '',
            '    # 3. SA主循环',
            '    T = T0',
            '    while T > T_min:',
            '        for _ in range(M0):',
            '            op = random.choice(["displace","add","delete"])',
            '            if op == "displace":',
            '                i = random.randint(0, N-1)',
            '                sigma_d = L * (T/T0)**0.5',
            '                parts_new = displace(particles, i, sigma_d)',
            '            elif op == "add":',
            '                parts_new = add_particle_maxent(particles)',
            '            else:',
            '                parts_new = delete_random_particle(particles)',
            '',
            '            # PAGCM评估 + Metropolis接受(迁移2: 统计物理)',
            '            m = PAGCM(parts_new, r0=250, alpha=0.5, L=10000)',
            '            m.solve_quiet()',
            '            Pc_new = sum(m.connectivity.values())/3',
            '            f_new = len(parts_new)/N_max + lambda_p*max(0,P_target-Pc_new)',
            '',
            '            delta_f = f_new - f_curr',
            '            if delta_f < 0 or random.random() < math.exp(-delta_f/T):',
            '                particles, N, f_curr = parts_new, len(parts_new), f_new',
            '                if f_curr < f_best: X_best, f_best = parts_new, f_curr',
            '',
            '        T *= gamma  # 指数冷却 (迁移3: 冶金退火)',
            '',
            '    return X_best, f_best',
        ]
        cap = q2.add_paragraph()
        r = cap.add_run('代码片段 Q2-1 MESA-PAGCM模拟退火优化核心代码 (完整代码见项目文档)')
        r.font.size = Pt(9); r.font.bold = True; r.font.italic = True
        target.addprevious(cap._element)
        for line in code_lines:
            cp = q2.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cp.add_run(line); r.font.size = Pt(8); r.font.name = 'Consolas'
            cp.space_after = Pt(2); cp.space_before = Pt(0)
            target.addprevious(cp._element)
        sp = q2.add_paragraph(); target.addprevious(sp._element)
        break

q2_v3 = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出\第二问_MESA-PAGCM完整报告_含图表_v3.docx'
q2.save(q2_v3)
print('[OK] Q2 v3: %.0f KB' % (os.path.getsize(q2_v3)/1024))
print('Done!')
