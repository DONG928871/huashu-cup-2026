# -*- coding: utf-8 -*-
"""使用python-docx生成第一问完整.docx，含文字+表格+嵌入PNG图片"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出'
IMG = os.path.join(OUT, '图片')
os.makedirs(OUT, exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_title(text, level=0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.space_before = Pt(18)
        p.space_after = Pt(10)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    return p

def add_body(text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.space_after = Pt(6)
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = 'Times New Roman'
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.bold = True
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()  # spacing
    return table

def add_image(filename, caption='', width_inches=5.5):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            run = p.add_run(caption)
            run.font.size = Pt(9)
            run.font.bold = True
        doc.add_paragraph()
        doc.add_picture(path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    else:
        add_body('[图片未找到: %s]' % filename)

# ===== BUILD DOCUMENT =====
add_title('A题 微构体中填充导电介质的仿真优化', 0)
add_title('第一问 完整建模报告', 0)
add_title('PAGCM — 周期边界自适应图连通判定模型', 0)
add_body('Periodic-Adaptive Graph Connectivity Model', indent=False)

# ═══ MODULE 1 ═══
add_title('模块一：模型建立与公式推导', 1)

add_title('1.1 变量定义三线表', 2)
add_body('表1-1 PAGCM模型变量定义三线表。按决策变量/中间变量/目标变量/已知参数/待校准参数分类，标注类型、单位、取值范围、现实场景含义。', indent=False)
add_table(
    ['变量符号','变量名称','变量类型','单位','取值范围','现实场景含义'],
    [
        ['p_i=(x_i,y_i,z_i)','粒子中心坐标','已知参数','坐标单位','[-5000,5000]^3','导电填料在RVE中的空间位置，由附件直接读取'],
        ['N','粒子总数','已知参数','个','{12,49,535}','微构体中导电填料粒子总数量'],
        ['L','RVE边长','已知参数','坐标单位','10000','代表体积单元立方体边长，从坐标极差推断'],
        ['r0','粒子基础几何半径','待校准参数','坐标单位','[50,500]','导电填料物理半径。附件未提供，需据题目场景合理假设'],
        ['alpha','自适应强度系数','待校准参数','无量纲','[0,2]','控制密度感知对等效半径的调节强度'],
        ['R_search','局部密度搜索半径','中间变量','坐标单位','1500','估算粒子i局部数密度时的球形搜索域半径'],
        ['rho_global','全局平均数密度','中间变量','L^-3','N/L^3','RVE内粒子平均空间密度，自适应基准'],
        ['rho_local(i)','粒子i局部数密度','中间变量','L^-3','[0,+inf)','搜索半径内近邻粒子数密度'],
        ['ri_eff','密度感知等效半径','中间变量','坐标单位','[0.5r0,3r0]','经局部密度修正的有效作用半径'],
        ['d_T(pi,pj)','环面距离','中间变量','坐标单位','[0,L*sqrt3/2]','周期边界下两粒子最短距离'],
        ['k_ij','周期偏移矢量','中间变量','Z^3','{-1,0,+1}^3','记录边(i,j)跨越周期边界的次数和方向'],
        ['G=(V,E)','周期图','决策变量','-','|V|=N,|E|可变','定义在3-环面T^3上的无向图'],
        ['C_k','第k个连通分量','决策变量','-','|Ck| in [1,N]','图的极大连通子图'],
        ['conn_dir','方向性连通判定','目标变量','布尔','{0,1}','判定该方向是否存在贯穿RVE的导电通路'],
        ['P_conn(dir)','连通概率','目标变量','-','[0,1]','MC扰动后统计的连通概率'],
    ])

add_title('1.2 核心模型假设', 2)
for name, content, basis, impact in [
    ('假设1: 粒子球形假设',
     '所有导电填料粒子为刚性球体，具有相同基础几何半径r0。导电接触仅依赖环面距离与等效半径之和的比较: d_T <= ri_eff+rj_eff。',
     '题目未给出粒子形状参数，球形是最简单的各向同性假设；逾渗理论经典模型(Balberg 1984, Scher & Zallen 1970)以球形粒子为基准。',
     '使距离判据简化为标量比较，避免取向相关判断。若非球形需引入形状因子修正。'),
    ('假设2: 局部密度线性响应假设',
     '等效半径对局部密度的响应是线性的: ri_eff = r0 * [1 + alpha * (rho_local/rho_global - 1)]。',
     '线性响应是最小假设(Occam剃刀)；一阶泰勒展开是光滑函数在偏离0附近的合理近似。',
     '若真实物理存在饱和效应，线性模型在极端密度比处产生偏差。通过alpha扫描和r_eff截断评估约束。'),
    ('假设3: 二值接触导电假设',
     '两粒子导电状态为二值开关: d_T <= ri_eff+rj_eff时导电桥形成(电阻=0)，否则断路(电阻无穷大)。',
     '逾渗理论经典框架基于二值连接；微米级填料接触导电远大于隧穿导电。',
     '忽略隧穿导电渐变特性。补偿: 可引入指数衰减连接概率p_ij=exp(-d_ij/xi_tunnel)。'),
    ('假设4: 周期边界统计均匀假设',
     'RVE是无限周期结构的代表性采样。边界处粒子分布与内部统计一致，无边界聚集效应。',
     'PBC是微力学RVE方法标准假设(Hill 1963)；附件坐标对称分布于[-5000,5000]且边界值频繁出现。',
     '消除边界条件选择的任意性。环面拓扑使连通判定仅取决于粒子分布本身。'),
    ('假设5: 静态几何模型假设',
     '粒子位置固定不变。忽略热运动、布朗运动、电迁移及基体固化中粒子重排。',
     '题目描述为确定性结构数据；聚合物固化后填料位置固定是工程常态。',
     '通过MC位置扰动(sigma=0.05r0)将确定论扩展为概率论框架补偿静态假设。'),
]:
    add_body('[%s] %s' % (name, content))
    add_body('设立依据: %s' % basis)
    add_body('对模型的影响: %s' % impact)

add_title('1.3 分步公式推导', 2)
add_body('以下推导严格遵循场景抽象-变量定义-模型构建-数据适配-求解迭代-结果验证六阶段逻辑链。全文符号统一无冲突。', indent=False)
for formula in [
    '(1) pi_prime = pi + (L/2, L/2, L/2), i=1,2,...,N  [坐标平移至[0,L]^3]',
    '(2) d_T(pi,pj) = sqrt(sum_{dim=1}^3 (min(|pi_dim-pj_dim|, L-|pi_dim-pj_dim|))^2)  [环面距离度量]',
    '(3) k_ij_dim = -1 if diff > L/2; +1 if diff < -L/2; 0 otherwise  [周期偏移矢量]',
    '(4) rho_global = N / L^3  [全局平均密度]',
    '(5) n_i = |{j != i : d_T(pi,pj) <= R_search}|  [局部近邻计数]',
    '(6) rho_local(i) = 3n_i / (4*pi*R_search^3)  [局部数密度]',
    '(7) ri_eff = r0 * [1 + alpha * (rho_local(i)/rho_global - 1)]  [密度感知等效半径—核心]',
    '(8) ri_eff = clamp(ri_eff, 0.5r0, 3.0r0)  [物理截断]',
    '(9) V = {1,2,...,N}  [节点集]',
    '(10) E = {(i,j,k_ij) : d_T(pi,pj) <= ri_eff+rj_eff, i<j}  [边生成规则]',
    '(11) parent[i]=i, rank[i]=0, for all i in V  [并查集初始化]',
    '(12) for each (i,j,k_ij) in E: Union(Find(i),Find(j))  [合并+路径压缩+按秩合并]',
    '(13) component[i]=relabel(Find(i)); C_k={i:component[i]=k}  [连通分量]',
    '(14) S_lo(d)={i: pi_dim-ri_eff<=0}  [低边界接触集]',
    '(15) S_hi(d)={i: pi_dim+ri_eff>=L}  [高边界接触集]',
    '(16) conn_d=1 if exists i in S_lo(d),j in S_hi(d): Find(i)=Find(j); else 0  [连通判定]',
    '(17) pi^(m)=pi+eps_i^(m), eps_i,d^(m)~N(0,sigma^2), sigma=0.05r0  [MC扰动]',
    '(18) P_conn(d)=(1/M)sum_{m=1}^M conn_d^(m), M=200  [连通概率]',
]:
    add_formula(formula)

add_title('1.4 建模流程图', 2)
add_image('图0_PAGCM建模流程图.png', '图0 PAGCM建模流程图（10步核心模块，四色阶段分组）')

# ═══ MODULE 2 ═══
add_title('模块二：模型求解与结果呈现', 1)

add_title('2.1 求解环境与步骤', 2)
add_body('求解语言：Python 3.11，纯标准库(zero external dependencies)。算法：网格空间索引+并查集(Union-Find)+BFS路径回溯。代码文件：q1_solve.py(约480行，逐行详细注释)。')
add_table(
    ['步骤','操作','输入','输出','复杂度'],
    [
        ['S1','数据加载','all_datasets.json','6组坐标列表','O(N)'],
        ['S2','坐标平移至[0,L]^3','原始坐标','归一化坐标','O(N)'],
        ['S3','局部密度估计','坐标,R_search','n_i列表','O(N^2)'],
        ['S4','自适应等效半径','n_i,alpha,r0','r_eff列表','O(N)'],
        ['S5','空间网格构建','归一化坐标','空间哈希表','O(N)'],
        ['S6','图边生成','r_eff,坐标','邻接边列表E','O(N*avg_deg)'],
        ['S7','并查集聚类','边列表E','component标签','O(|E|*alpha(N))'],
        ['S8','方向连通判定','component,r_eff','conn_X/Y/Z','O(|S_lo|*|S_hi|)'],
        ['S9','BFS路径回溯','源/目标粒子','最短导通路径','O(|E|)'],
        ['S10','MC扰动(M=200)','sigma=12.5','P_conn','O(M*N*avg_deg)'],
    ])

add_title('2.2 求解结果', 2)
add_body('表2-1 PAGCM连通性判定结果汇总(r0=250, alpha=0.5, L=10000)', indent=False)
add_table(
    ['数据集','N','X方向','Y方向','Z方向','边数','分量数','最大簇','r_eff均值'],
    [
        ['组1_场景A','12','不连通','不连通','不连通','15','1','12','750.0'],
        ['组1_场景B','12','不连通','不连通','不连通','16','3','7','697.9'],
        ['组2_场景A','49','连通','不连通','不连通','321','1','49','750.0'],
        ['组2_场景B','49','不连通','不连通','不连通','306','1','49','750.0'],
        ['组3_场景A','535','连通','连通','连通','198','384','30','269.4'],
        ['组3_场景B','535','连通','连通','连通','173','401','19','269.7'],
    ])
add_body('表2-2 蒙特卡洛连通概率(M=200, sigma=12.5)', indent=False)
add_table(
    ['数据集','P_conn(X)','P_conn(Y)','P_conn(Z)','结论'],
    [
        ['组1_场景A','0.000','0.000','0.000','全方向绝缘—确定性结论'],
        ['组1_场景B','0.000','0.000','0.000','全方向绝缘—确定性结论'],
        ['组3_场景A','0.755','1.000','1.000','X方向逾渗临界区(敏感); Y/Z方向鲁棒'],
    ])

add_title('2.3 结果可视化图表', 2)
add_image('图1_连通性热力图.png', '图1 连通性判定热力图（红色=连通，蓝色=不连通）')
add_image('图2_MC连通概率.png', '图2 蒙特卡洛连通概率（M=200轮，sigma=12.5）')
add_image('图3_PAGCM_vs_GPNM对比.png', '图3 PAGCM vs GPNM连通方向数对比')
add_image('图4_alpha敏感度曲线.png', '图4 alpha参数敏感度扫描(组3_场景A)')
add_image('图5_连通分量统计.png', '图5 连通分量统计与等效半径分布')
add_image('图6_性能与验证.png', '图6 求解性能与PAGCM vs GPNM交叉验证')

# ═══ MODULE 3 ═══
add_title('模块三：模型检验与验证', 1)

add_title('3.1 有效性检验——PAGCM vs GPNM交叉验证', 2)
add_body('检验方法：对全部6数据集x3方向=18个判定，分别运行PAGCM(alpha=0.5)和GPNM(固定r0=250)，逐项对比。GPNM使用O(N^2)暴力计算确保零近似误差。')
add_body('检验结果：16/18=88.9%判定完全一致。2处差异(组2A_X和组3A_X)经手动距离验证确认PAGCM正确——自适应半径在稀疏区将等效半径从250扩大到约750，捕捉到GPNM遗漏的真实逾渗路径。PAGCM检出率比GPNM提升11.1%。')

add_title('3.2 鲁棒性分析——参数+/-10%到+/-20%波动', 2)
add_body('对组1_A(远低于临界值)和组3_A(临界区)施加alpha+/-20%和r0+/-10%波动。')
add_table(
    ['数据集','参数','波动','取值','变化','稳健性','理由'],
    [
        ['组1_A','alpha','+/-20%','0.4/0.6','无变化','极稳健','粒子极稀疏'],
        ['组1_A','r0','+/-10%','225/275','无变化','极稳健','12粒子不足贯穿'],
        ['组3_A','alpha','-20%','0.4','X:通->断','敏感','临界区特征'],
        ['组3_A','alpha','+20%','0.6','无变化','稳健','Y/Z方向可靠'],
        ['组3_A','r0','-10%','225','X:通->断','敏感','临界区特征'],
        ['组3_A','r0','+10%','275','无变化','稳健','扩大r0增强连接'],
    ])
add_body('结论：组3_X方向在alpha<=0.4或r0<=225时连通性翻转——这不是模型缺陷，而是该方向处于逾渗临界区的物理本质。临界点附近连通性对参数高度敏感(关联长度发散，临界指数beta约0.4)。工程意义：若需X方向可靠导电，建议增加约10%填料量(535->590)远离临界区。')

add_title('3.3 模型对比——PAGCM vs GPNM六维度量化', 2)
add_table(
    ['维度','GPNM(固定半径)','PAGCM(自适应半径)','PAGCM改进'],
    [
        ['阈值策略','d_T<=2r0固定','d_T<=ri_eff+rj_eff','消除人为偏差'],
        ['检出灵敏度','仅几何接触','几何+近场耦合','检出率+11.1%'],
        ['可标定性','无自由参数','alpha可标定至实验','物理可校准'],
        ['计算效率','O(N^2)~0.15s','网格加速~0.20s','效率相近'],
        ['输出信息','1bit/方向','二值+概率+分量+路径','信息量x4'],
        ['可解释性','纯几何判据','密度感知+逾渗理论','可解释临界行为'],
    ])

# ═══ MODULE 4 ═══
add_title('模块四：结果深度分析与讨论', 1)

add_title('4.1 基础数值分析', 2)
add_body('(1) 组1场景(12粒子, phi约0.08%)全方向绝缘——需增加至>=200粒子(phi约1.3%)或改用高长径比填料(碳纳米管/石墨烯)以降低逾渗阈值。')
add_body('(2) 组2场景A(49粒子, phi约0.32%)仅X方向单向导电——若工程仅需单向导电已满足；若需三维导电需增加至>=200或通过链状排列引导YZ方向取向(不增加填料量的免费优化)。')
add_body('(3) 组3场景(535粒子, phi约3.5%)三维导电已实现——但X方向P_conn仅75.5%提示安全裕度不足。建议增加10%填料量至约590粒子使X方向P_conn>=99%。')

add_title('4.2 深层机理分析', 2)
add_body('逾渗理论三阶段验证：(a)亚临界区(组1)——孤立团簇，无贯穿路径。(b)逾渗转变区(组2)——方向性逾渗开始出现，各向异性显著。(c)超临界区(组3)——三方贯穿但临界区特征显著(384分量/最大簇30，分形结构)。Y/Z方向100%鲁棒说明网络有冗余路径；X方向75.5%敏感说明该方向仅1-2条瓶颈路径。')
add_body('约束条件验证：全部粒子在RVE内(空间约束满足)，环面度量自动满足周期边界等价关系(拓扑约束满足)，384<535且最大簇30>贯穿路径长5(物理合理性满足)。')

add_title('4.3 应用延伸分析', 2)
add_body('工程建议三级：(1)组1场景需大幅改进——增加约16倍填料量或采用纳米填料。(2)组2场景可微调优化——通过剪切流动或电场取向使粒子沿YZ方向链化排列实现免费优化。(3)组3场景需小幅微调——增加10%填料量提升X方向可靠性至>=99%。')
add_body('模型推广：PAGCM的方法论(空间索引+密度感知自适应+周期拓扑嵌入+并查集聚类)可推广至热导率逾渗、力学增强网络、多孔介质渗流、电化学离子传输网络等涉及连通性判定的广泛工程问题。三维环面上的图连通模型为周期性复合材料微观-宏观均质化提供了离散、计算可操作的桥接工具。')

# ===== SAVE =====
path = os.path.join(OUT, '第一问_PAGCM完整报告_含图表.docx')
doc.save(path)
print('[OK] %s  (%.0f KB)' % (path, os.path.getsize(path)/1024))
