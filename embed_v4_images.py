# -*- coding: utf-8 -*-
"""将8张图片嵌入完整论文v4的合适位置"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document(r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\完整论文_v4.docx')
img_dir = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\论文输出\图片'

# Remove old missing file warnings
to_remove = []
for i, p in enumerate(doc.paragraphs):
    if '图片文件未找到' in p.text:
        to_remove.append(i)
for i in reversed(to_remove):
    p_element = doc.paragraphs[i]._element
    p_element.getparent().remove(p_element)
print(f'Removed {len(to_remove)} missing file warnings')

def embed(target_p, img_file, caption, width=5.0):
    img_path = os.path.join(img_dir, img_file)
    if not os.path.exists(img_path):
        print(f'  [SKIP] {img_file} not found')
        return
    el = target_p._element
    # caption
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.font.size = Pt(9); r.font.bold = True
    el.addprevious(cp._element)
    # image
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(img_path, width=Inches(width))
    el.addprevious(ip._element)
    # spacing
    sp = doc.add_paragraph(); el.addprevious(sp._element)
    print(f'  [OK] {img_file}')

# Use exact paragraph indices from the document analysis
# Key insertion points:
# P53: "模型架构如图1所示" -> Fig1
# P92: "图2展示49粒子路径" -> Fig7
# P93: "图3" for 535粒子 -> Fig8
# P102: "如图1所示" in MC-GCPM section -> Fig2
# P103: missing file warning for fig_q2 -> Fig2
# P125: Table 2 MC results -> Fig5 after it
# P129: "图1展示连通概率随填充率变化" -> Fig2 (remove text ref)
# P6 (abstract): mentions 递进 -> Fig4 after problem analysis section
# P7 (abstract): mentions 双因素成本优化 -> Fig6 after Q4 section

# Map: paragraph index -> (img_file, caption, width)
# Using indices AFTER the text that references the figure
insert_points = {
    53: ('v4_fig1_gcpm_flowchart.png', '图1 GCPM模型分析流程图（三层架构：几何层KD-Tree+GMP -> 过滤层并查集 -> 输出层连通判定+BFS路径回溯）', 5.0),
    92: ('v4_fig7_49path.png', '图2 49粒子最短导通路径示意（路径：#5(左电极) -> #14 -> #33 -> #2(右电极)，共4步，52对电极间连通）', 4.5),
    93: ('v4_fig8_535network.png', '图3 535粒子大规模导电网结构示意（189粒子骨架/35.3%，路径含62节点）', 4.5),
    102: ('v4_fig2_percolation_curve.png', '图4 MC-GCPM渗透转变曲线（500次MC/组，95% Clopper-Pearson置信区间，Logistic S曲线拟合）', 5.0),
    125: ('v4_fig5_mc_histogram.png', '图5 500次MC实验连通频数分布（4组填充率phi=0.50%/0.60%/0.70%/1.00%对比）', 5.5),
    129: ('v4_fig3_binary_search.png', '图6 自适应二分搜索算法流程图（Clopper-Pearson CI判断+收敛判据，输出phi_min=0.82%, N=580）', 5.0),
    23: ('v4_fig4_progression.png', '图7 四问递进关系与方法链（确定性连通判定 -> 不确定统计分析 -> 临界点精确定位 -> 双因素成本优化）', 5.5),
    31: ('v4_fig6_cost_comparison.png', '图8 双因素成本优化——A/B介质成本与导通效率对比（A效率约为B的44倍）', 5.0),
}

# Insert in reverse index order to avoid shifting
for idx in sorted(insert_points.keys(), reverse=True):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        fn, cap, w = insert_points[idx]
        embed(p, fn, cap, w)

out_path = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\论文输出\完整论文_v4_含图表.docx'
doc.save(out_path)
print(f'\n[OK] {out_path} ({os.path.getsize(out_path)/1024:.0f} KB)')
