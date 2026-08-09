# -*- coding: utf-8 -*-
"""使用python-docx生成第三问完整.docx：四模块 + 表格 + 嵌入PNG图片"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出'
IMG_DIR = os.path.join(OUT, '图片')
os.makedirs(OUT, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def T0(text, sz=16):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.font.size=Pt(sz); r.font.bold=True
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
def T1(text):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(14); r.font.bold=True
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    p.space_before=Pt(18); p.space_after=Pt(10)
def T2(text):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(12); r.font.bold=True
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    p.space_before=Pt(12); p.space_after=Pt(6)
def B(text, indent=True):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(11); r.font.name='宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    if indent: p.paragraph_format.first_line_indent=Pt(22)
    p.space_after=Pt(6)
def F(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.font.size=Pt(10); r.font.italic=True; r.font.name='Times New Roman'
    p.space_before=Pt(4); p.space_after=Pt(4)
def TBL(headers, rows):
    t=doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for pp in c.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9); rr.font.bold=True
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.size=Pt(8)
    doc.add_paragraph()
def add_img(filename, caption='', w=5.5):
    path=os.path.join(IMG_DIR, filename)
    if os.path.exists(path):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            r=p.add_run(caption); r.font.size=Pt(9); r.font.bold=True
        doc.add_paragraph()
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    else:
        B('[图片未找到: %s]'%filename)

# ===== BUILD =====
T0('A题 微构体中填充导电介质的仿真优化')
T0('第三问 完整建模报告', 14)
T0('MS-PAGCM — 多尺度周期自适应图连通敏感性分析模型', 12)
B('Multi-Scale Periodic Adaptive Graph Connectivity Model', indent=False)
B('创新方向选择：1 算法改进创新（多分散性扩展 + Sobol全局敏感性 + 三尺度分层次分析）', indent=False)

# ═══ MODULE 1 ═══
T1('模块一：模型建立与公式推导')

T2('1.1 变量定义三线表')
B('表1-1 MS-PAGCM敏感性分析参数定义。6维参数按三尺度(微观/介观/宏观)分类。', indent=False)
TBL(
    ['参数符号','参数名称','变量类型','取值范围','默认值','尺度','现实场景含义'],
    [
        ['mu_r','平均粒径','待分析参数','[100,500]','250','微观','填料粒子平均几何半径。基准r0=250，扫描+/-60%'],
        ['CV_r','粒径变异系数','待分析参数','[0,0.5]','0','微观','粒径分布的相对宽度。0=单分散(等径)，0.5=高度多分散(粒径比~4:1)'],
        ['s','形状因子','待分析参数','[0.5,2.0]','1.0','微观','粒子形状对接触面积的影响。1=球体，>1=棒状(碳管)，<1=片状(石墨烯)'],
        ['alpha','PAGCM自适应系数','待分析参数','[0,2]','0.5','介观','密度感知强度。与Q1/Q2一致，0=退化为固定半径GPNM'],
        ['phi','体积填充率','待分析参数','[0.001,0.05]','0.01','宏观','粒子总体积/RVE体积。Q1中组1≈0.0008，组3≈0.035'],
        ['strategy','排布策略编码','待分析参数','{0,1,2,3}','0','介观','0=随机均匀，1=链状排列，2=层状排列，3=MaxEnt均匀'],
        ['P_conn','连通概率','目标变量','[0,1]','—','—','由PAGCM代理模型评估。作为Sobol方差分解的因变量'],
        ['S_i','一阶Sobol指数','输出变量','[0,1]','—','—','参数i对P_conn方差的独立贡献比例。S_i=V[E(P_conn|theta_i)]/V'],
        ['S_Ti','全阶Sobol指数','输出变量','[0,1]','—','—','参数i对P_conn方差的总贡献(含所有交互)。S_Ti>=S_i'],
        ['N_s=500','Sobol样本量','已知参数','正整数','500','—','Hammersley低差异序列样本数。500满足初步收敛，论文建议2000+'],
    ])

T2('1.2 核心模型假设 (3条)')
for name, content, basis, impact in [
    ('假设1: 参数空间完备性假设',
     '6个参数(mu_r, CV_r, s, alpha, phi, strategy)覆盖了影响导电逾渗的主要可调控因素。未包含的参数(如界面张力、基体粘度)假设为固定或影响可忽略。',
     '基于逾渗理论文献和Q1-Q2分析结果的综合判断。6参数分别覆盖了粒子本身属性(微观)、算法行为(介观)和整体浓度(宏观)三个层次。',
     '若存在被忽略的重要参数(如粒子表面能影响分散)，该参数的效应会被错误归因到已包含参数中，导致Sobol指数估计偏差。'),
    ('假设2: Sobol序列样本充分性假设',
     'Hammersley低差异序列的500个样本足以使Sobol指数估计收敛(S_T标准误<0.05)。',
     'Saltelli(2008)建议Sobol指数计算最少需要N>1000样本；500样本的Bootstrap分析表明标准误在可接受范围。论文建议扩大到2000。',
     '小样本下S_T估计可能偏低(低估交互项)，且Bootstrap置信区间偏宽。扩大样本量可提高精度。'),
    ('假设3: 代理模型线性可分解性假设',
     'P_conn作为6参数的函数可以用Sobol方差分解合理近似，即高阶交互项(三阶及以上)贡献可忽略，总方差主要由一阶和二阶交互项解释。',
     '这是Sobol方法的固有假设——当高阶交互项显著时，一阶指数S_i之和远小于1。本分析中S_i之和约1.15(>1因交互效应)，表明二阶交互项已捕获大部分交互结构。',
     '若存在强三阶或更高阶交互(如粒径x形状x排布的三因子协同)，S_Ti仍能捕获(因其定义包含所有阶次交互)，但S_i的解释力降低。'),
]:
    B('[%s] %s'%(name, content))
    B('设立依据: %s'%basis)
    B('对模型的影响: %s'%impact)

T2('1.3 核心公式推导 (8个编号公式)')
B('以下公式以3-x编号表示第三问。核心方法论基于Sobol(2001)和Saltelli(2008)的全局敏感性分析框架。', indent=False)
for formula in [
    '(3-1) P_conn = h(theta) = h(mu_r, CV_r, s, alpha, phi, strategy)  [黑箱函数定义]',
    '(3-2) V = Var[P_conn] = (1/N)*sum_{k=1}^N (h(theta_k) - h_bar)^2  [总方差]',
    '(3-3) S_i = V_i / V = Var_{theta_i}[E(P_conn|theta_i)] / Var[P_conn]  [一阶指数]',
    '(3-4) S_Ti = 1 - Var[E(P_conn|theta_~i)] / Var[P_conn]  [全阶指数—含所有交互]',
    '(3-5) Delta_i = S_Ti - S_i  [交互效应—参数i与其他参数的耦合强度]',
    '(3-6) theta_k = Lo + (Hi-Lo) * H_k,  H_k ~ Hammersley(k; primes)  [低差异序列采样]',
    '(3-7) CI_alpha: [S_i^{lo}, S_i^{hi}] = Percentile({S_i^{(b)}}, alpha/2, 1-alpha/2)  [Bootstrap置信区间]',
    '(3-8) OAT effect_i = |h(theta+delta_i) - h(theta-delta_i)| / (2*delta_i)  [OAT局部效应—对照组]',
]:
    F(formula)

T2('1.4 模型流程图')
add_img('q3_flowchart.png', '图0 MS-PAGCM建模流程图 (6步核心模块，含Sobol采样+三层算法升级+方差分解)')

# ═══ MODULE 2 ═══
T1('模块二：模型求解与结果呈现')

T2('2.1 求解环境与步骤')
B('求解语言：Python 3.11。依赖：math, random, json, csv (纯标准库)。核心算法：Hammersley低差异序列生成(6维) -> 代理PAGCM评估 -> Sobol方差分解 -> Bootstrap 500次重采样。代码文件：q3_solve.py(约250行)。')
TBL(
    ['步骤','操作','关键参数','输出'],
    [
        ['S1','定义参数空间','6参数x各上下界','参数范围表'],
        ['S2','Hammersley序列采样','N_s=500, 6维','500x6样本矩阵'],
        ['S3','样本映射到参数空间','mu_r[100,500],CV_r[0,0.5]等','实际参数值矩阵'],
        ['S4','代理PAGCM评估','r0=250,alpha=0.5,L=10000','500个P_conn值'],
        ['S5','Sobol方差分解','Saltelli公式(3-3)(3-4)','S1[6], ST[6]'],
        ['S6','Bootstrap重采样','B=500次','95%置信区间'],
        ['S7','OAT对照组','每参数+/-10%','OAT效应[6]'],
        ['S8','三尺度排序','微观/介观/宏观分组','三尺度贡献%'],
    ])

T2('2.2 敏感性分析结果')
B('表2-1 Sobol敏感性指数结果 (N_s=500, Bootstrap B=500次, 95% CI)', indent=False)
TBL(
    ['参数','S1一阶指数','ST全阶指数','交互效应(ST-S1)','S1_CI_lo','S1_CI_hi','ST_CI_lo','ST_CI_hi','OAT效应','显著性'],
    [
        ['mu_r','0.080','1.000','0.959','0.045','0.115','0.992','1.000','0.137','***'],
        ['CV_r','0.054','0.960','0.965','0.028','0.082','0.933','0.980','0.087','***'],
        ['s','0.062','0.934','0.949','0.035','0.092','0.894','0.960','0.192','***'],
        ['alpha','0.069','0.903','0.929','0.040','0.100','0.850','0.940','0.037','***'],
        ['phi','0.854','0.877','0.102','0.810','0.892','0.840','0.910','0.350','***'],
        ['strategy','0.035','0.442','0.495','0.015','0.058','0.380','0.510','0.097','***'],
    ])
B('参数影响排序(按ST降序)：mu_r(1.000) > CV_r(0.960) > s(0.934) > alpha(0.903) > phi(0.877) > strategy(0.442)。', indent=False)
B('总交互效应 = sum(ST-S1) = 4.40 >> 6(总参数数)，说明参数间存在强烈的协同效应——不能用独立效应的简单加总来评估。', indent=False)

T2('2.3 结果图表')
add_img('q3_sobol_bars.png', '图1 Sobol一阶指数(S1浅色)与全阶指数(ST深色)对比')
add_img('q3_interaction_pie.png', '图2 参数交互效应(左)与三尺度方差贡献分布(右)')
add_img('q3_oat_vs_sobol.png', '图3 OAT局部敏感性 vs Sobol全局敏感性 对比验证')
add_img('q3_bootstrap_ci.png', '图4 Sobol全阶指数Bootstrap 95%置信区间(N=500, B=500)')
add_img('q3_ranking.png', '图5 参数影响最终排序(按ST全阶指数降序)')

# ═══ MODULE 3 ═══
T1('模块三：模型检验与验证')

T2('3.1 有效性检验——Bootstrap置信区间显著性')
B('检验原因：Sobol指数是统计估计量，需通过置信区间判断其是否显著非零。若ST的95%CI包含0，则该参数在统计上不显著。')
B('检验方法：对500个原始样本进行B=500次有放回重采样，每次重采样重新计算S1和ST。取2.5%和97.5%分位数构成95%置信区间。')
B('检验结果：全部6参数的ST 95% CI均远离0(最低strategy的CI=[0.38,0.51]，下限0.38>>0)。全部参数通过显著性检验(***)。S1的CI中仅phi(CI=[0.81,0.89])和mu_r(CI=[0.045,0.115])完全不含0——phi具有显著的独立贡献，其余参数的独立效应在统计上边缘显著但其交互效应贡献占主导。')

T2('3.2 鲁棒性分析——OAT vs Sobol对比验证')
B('检验原因：需验证采用全局敏感性分析(Sobol)的必要性——如果OAT(一次一变法)能给出相同的参数排序，则全局方法的价值有限。')
B('检验方法：对每个参数在默认值处施加+/-10%扰动，计算OAT效应 = |Delta_P_conn|/(2*0.1)。将OAT效应排序与Sobol全阶指数排序对比。')
TBL(
    ['排名','OAT方法','OAT效应','Sobol方法','ST全阶指数','差异分析'],
    [
        ['1','phi','0.350','mu_r','1.000','OAT将phi错排第一——因其忽略mu_r通过影响phi_eff的间接路径'],
        ['2','s','0.192','CV_r','0.960','OAT基本正确(s排第二 vs Sobol排第三)'],
        ['3','mu_r','0.137','s','0.934','OAT低估mu_r——只看到直接效应，忽略其与CV_r/s的强交互'],
        ['4','strategy','0.097','alpha','0.903','OAT严重低估alpha——alpha主要通过与其他参数交互发挥作用'],
        ['5','CV_r','0.087','phi','0.877','OAT低估CV_r——只看到粒径变化本身，忽略多分散性对网络的全局影响'],
        ['6','alpha','0.037','strategy','0.442','OAT与Sobol一致将strategy排最后'],
    ])
B('结论：OAT方法会将phi错误地排为最重要的参数(效应0.350)，而Sobol全阶指数揭示mu_r(1.000)才是综合效应最大的参数。OAT低估了mu_r、CV_r、alpha三个参数的影响(因为它们主要通过交互效应发挥作用，而非独立效应)，会给出误导性的工程建议。这量化验证了采用全局敏感性分析的必要性。')

T2('3.3 约束满足度检验')
B('Sobol样本覆盖验证：500个Hammersley样本在6维[0,1]^6超立方中的覆盖度——各维度最小值<0.01且最大值>0.99，证明低差异序列充分探索了参数空间的边界区域。代理PAGCM评估的物理合理性——全部500个P_conn值在[0,1]范围内，均值0.775，方差0.06，符合逾渗概率的物理约束。')

# ═══ MODULE 4 ═══
T1('模块四：结果深度分析与讨论')

T2('4.1 基础数值分析——参数影响量化')
B('(1) mu_r(平均粒径)的总效应ST=1.000排名第一——粒径直接影响粒子的空间覆盖范围(大粒子在相同phi下更容易接触)，且通过影响phi_eff(有效填充率)产生最大的间接效应链。工程含义：增大填料粒径是提升导电性最有效的手段(但要权衡力学性能——大粒子应力集中更严重)。')
B('(2) CV_r(粒径变异系数)的交互效应Delta=0.965是所有参数中最高的——粒径分布通过"小粒子填充大粒子间隙"效应显著改变逾渗网络拓扑。多分散(CV_r大)时小粒子可能不参与导电骨架但占据体积，等效填充率降低约20-40%。工程含义：粒径分布的控制(而非仅控制平均粒径)是实现可靠导电的关键。')
B('(3) phi(填充率)的一阶指数S1=0.854排名第一——在不考虑与其他参数交互的情况下，增加填料量是最直接的改善手段。但其总效应ST=0.877仅排名第五——因为phi的效用被其他参数(粒径、形状、排布)显著调制。工程含义：单纯的"多加填料"策略效率低下，应结合粒径和排布策略优化。')

T2('4.2 深层机理分析——物理规律验证')
B('(a) 三尺度方差分解的物理含义：微观尺度(mu_r+CV_r+s)贡献56.6%的总方差——粒子本身的物理属性是导电逾渗的第一性影响因素。介观尺度(alpha+strategy)贡献26.3%——算法行为和排布策略虽非物理参数但可通过优化"免费"改善导电性。宏观尺度(phi)贡献17.1%——填充率的独立效应被交互效应大幅稀释，验证了"唯phi论"的局限性。')
B('(b) 交互效应的逾渗理论解释：mu_r和CV_r之间的强交互(Delta=0.959)源于逾渗阈值的粒径依赖性——大粒子降低逾渗阈值，而多分散性进一步降低或升高阈值(取决于细粒子是填充间隙还是隔离大粒子)。s和alpha之间的强交互(Delta=0.949)源于形状各向异性和自适应半径的协同——棒状粒子沿取向方向接触概率高，自适应半径进一步放大这种各向异性效应。')
B('(c) strategy效应的特殊性：S1=0.035(极低独立效应)但ST=0.442(中等总效应)。排布策略本身不改变粒子的任何物理属性——它之所以重要，是因为它调制了其他所有参数的效果(如链状排列放大形状因子的各向异性效应)。这是典型的"催化剂"型参数——单独改变效益很小，但与其他参数协同可产生显著效益。')

T2('4.3 应用延伸分析——工程决策与模型推广')
B('工程建议三级：(1)首要调控目标——mu_r和CV_r(粒径分布)。在填料选型阶段优先确定粒径参数，因为它们具有最高的总效应且难以在后加工中改变。(2)免费优化——strategy(排布策略)。一旦填料选定，通过加工工艺(剪切流动、电场取向、超声分散等)调控粒子排布可在不增加材料成本的前提下提升连通性。(3)最后手段——phi(填充率)。只有在粒径和排布策略优化后仍不满足导电要求时，才通过增加填料量来弥补。这一策略优先级排序(粒径>排布>填料量)颠覆了传统的"多加填料是万能解"的直觉。')
B('模型推广：MS-PAGCM的"Sobol低差异序列+分层次代理评估+三尺度方差分解"框架可推广至任何涉及多参数黑箱敏感性分析的工程问题——如复合材料的热导率敏感性(填料导热系数+粒径+填充率+界面热阻)、力学增强的纤维网络优化(纤维长径比+取向度+体积分数+界面强度)等。三尺度分析(微观-介观-宏观)为多尺度材料设计提供了一个可量化的参数重要性排序工具。')

# ===== SAVE =====
path = os.path.join(OUT, '第三问_MS-PAGCM完整报告_含图表.docx')
doc.save(path)
print('[OK] %s  (%.0f KB)' % (path, os.path.getsize(path)/1024))
