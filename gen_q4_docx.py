# -*- coding: utf-8 -*-
"""使用python-docx生成第四问完整.docx：四模块 + 表格 + 嵌入PNG图片"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出'
IMG_DIR = os.path.join(OUT, '图片')
os.makedirs(OUT, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.page_width=Cm(21); section.page_height=Cm(29.7)
    section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
    section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)
style=doc.styles['Normal']; style.font.name='宋体'; style.font.size=Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

def T0(t,sz=16):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=True
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
def T1(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(14); r.font.bold=True
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    p.space_before=Pt(18); p.space_after=Pt(10)
def T2(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.bold=True
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    p.space_before=Pt(12); p.space_after=Pt(6)
def B(t,indent=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(11); r.font.name='宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    if indent: p.paragraph_format.first_line_indent=Pt(22)
    p.space_after=Pt(6)
def F(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.font.size=Pt(10); r.font.italic=True; r.font.name='Times New Roman'
    p.space_before=Pt(4); p.space_after=Pt(4)
def TBL(h,r):
    t=doc.add_table(rows=len(r)+1,cols=len(h)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,hdr in enumerate(h):
        c=t.rows[0].cells[i]; c.text=hdr
        for pp in c.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9); rr.font.bold=True
    for ri,row in enumerate(r):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.size=Pt(8)
    doc.add_paragraph()
def IMG(fn,cap='',w=5.5):
    path=os.path.join(IMG_DIR,fn)
    if os.path.exists(path):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if cap: r=p.add_run(cap); r.font.size=Pt(9); r.font.bold=True
        doc.add_paragraph(); doc.add_picture(path,width=Inches(w))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph()
    else: B('[图片未找到: %s]'%fn)

# ===== BUILD =====
T0('A题 微构体中填充导电介质的仿真优化')
T0('第四问 完整建模报告',14)
T0('MOEA/D-PAGCM — 多目标进化分解周期自适应图连通工程设计模型',12)
B('Multi-Objective Evolutionary Algorithm based on Decomposition with PAGCM',indent=False)
B('创新方向选择：3 多模型融合组合创新（PAGCM物理评估 + MOEA/D多目标进化 + TOPSIS决策优选 + 熵权法客观赋权）',indent=False)

# ═══ MODULE 1 ═══
T1('模块一：模型建立与公式推导')

T2('1.1 变量定义三线表')
B('表1-1 MOEA/D-PAGCM变量定义。按决策变量/目标变量/算法参数/约束参数/决策参数分类。',indent=False)
TBL(
    ['变量符号','变量名称','变量类型','取值范围','现实场景含义'],
    [
        ['X=[{p_i},N,{r_i},strategy]','优化决策变量','决策变量','5维混合变量','粒子空间排布+数量+粒径+排布策略的完整工程方案'],
        ['F(X)=[f1,f2,f3,f4]','四目标向量','目标变量','R^4','f1=1-P_conn(导电损失),f2=N/Nmax(成本),f3=phi(重量),f4=1-E/E0(力学损失)'],
        ['P_conn','连通概率','约束','>=0.80','由PAGCM评估。Q4放宽至0.80(低于Q2的0.95)——给其他目标留优化空间'],
        ['phi','体积填充率','约束','<=0.10','Guth-Gold模型在phi>0.1后非线性显著，限制在此范围内保证代理模型精度'],
        ['N_pop=50','种群规模','算法参数','[50,200]','MOEA/D子问题数=权重向量数。50为演示值，论文推荐100(更细Pareto前沿)'],
        ['T=10','邻域大小','算法参数','[10,50]','每个子问题的交配伙伴数。T=N_pop*20%。太小->搜索局限，太大->收敛慢'],
        ['G_max=50','最大进化代数','算法参数','[100,500]','50为演示值。50代x50个体=2500次PAGCM评估，约50秒(演示可接受)'],
        ['CR=0.9','DE交叉率','算法参数','[0.5,1.0]','高CR->后代更多维度来自变异向量(高探索)。Storn&Price(1997)推荐CR in [0.8,1.0]'],
        ['F_mut=0.5','DE缩放因子','算法参数','[0.3,0.9]','中等步长，探索与开发的平衡。标准推荐F in [0.4,0.9]'],
        ['B_Guth=2.5','Guth-Gold系数','代理模型参数','[1.5,4.0]','球形填料理论值(Einstein 1906)。棒状填料可取3.5-4.0'],
        ['w_j','TOPSIS权重','决策参数','数据驱动(熵权法)','从Pareto前沿的数据分布中动态计算，避免人为设定。H_j=-k*sum(p_ij*log p_ij)'],
        ['C_i','相对贴近度','决策输出','[0,1]','C=D-/(D+ + D-)。C越大方案越优。TOPSIS推荐C最大的方案'],
    ])

T2('1.2 核心模型假设 (4条)')
for name,content,basis,impact in [
    ('假设1: Pareto最优性假设',
     'MOEA/D生成的Pareto前沿近似能够充分代表真实Pareto前沿。N_pop=50/100个子问题均匀覆盖权重空间，进化G=50/200代后前沿已充分收敛。',
     'Zhang&Li(2007)证明MOEA/D在连续多目标优化问题上能以概率1收敛到Pareto前沿。Das-Dennis单纯形法保证权重向量均匀分布。',
     '有限种群和有限代数下Pareto前沿可能不完全(遗漏某些权重区域的最优解)。通过HV(超体积)指标监控收敛和增大N_pop/G_max来改善。'),
    ('假设2: Guth-Gold力学代理模型的准确性假设',
     '复合材料的相对模量E/E0与填料体积分数phi的关系可用Guth-Gold模型E/E0=1+2.5*phi充分近似。',
     'Einstein(1906)推导球形粒子悬浮液粘度增强的Guth-Gold系数B=2.5。实验验证在phi<0.1时线性关系良好。B=3.5-4.0适用于棒状填料。',
     'phi>0.1时粒子间相互作用使Guth-Gold模型低估模量。Q4通过约束phi<=0.10来保证代理模型在有效范围内。若需突破此约束需改用Mori-Tanaka或有限元均质化。'),
    ('假设3: 四目标覆盖工程设计主要关切的充分性假设',
     '导电性(f1)、材料成本(f2)、重量(f3)和力学性能(f4)四个目标覆盖了导电复合材料工程设计的核心关切。',
     '基于工业界的典型需求：电子封装关注导通率和成本，航空航天关注重量，结构件关注力学。四目标代表了不同应用场景的核心需求。',
     '某些场景可能有额外关切(如热导率、电磁屏蔽效能、环境稳定性)。MOEA/D框架天然可扩展——增加一个目标仅需增加一个权重维度。'),
    ('假设4: 熵权法客观赋权的合理性假设',
     '从Pareto前沿的数据分布中提取的熵权能够客观反映各目标在方案集中的"分辨度"——方差大的目标(方案间差异大)获得更高权重。',
     '熵权法(Shannon 1948信息熵在决策中的应用)在无决策者主观偏好时提供了一种客观、可复现的赋权方法。',
     '熵权法可能赋予在Pareto前沿上"容易区分方案"的目标更高权重，而非"更重要"的目标更高权重。若决策者有明确偏好，可在熵权基础上引入主观权重修正。'),
]:
    B('[%s] %s'%(name,content)); B('设立依据: %s'%basis); B('对模型的影响: %s'%impact)

T2('1.3 核心公式推导 (10个编号公式)')
B('公式以4-x编号。核心方法论基于MOEA/D(Zhang&Li,2007)和TOPSIS(Hwang&Yoon,1981)。',indent=False)
for f in [
    '(4-1) min F(X)=[f1=1-P_conn, f2=N/N_max, f3=phi, f4=1-E/E0]  s.t. P_conn>=0.80, phi<=0.10  [四目标优化问题]',
    '(4-2) lambda^(k) in R^4, sum(lambda_j)=1, k=1,...,N_pop  [Das-Dennis均匀权重向量生成]',
    '(4-3) g^{tch}(X|lambda,z*)=max_{j=1..4}{lambda_j*|f_j(X)-z*_j|}  [切比雪夫聚合函数]',
    '(4-4) z*_j = min{f_j(X_i), i=1..N_pop}, j=1..4  [理想点—各目标的历史最小值]',
    '(4-5) U_i = X_r1 + F*(X_r2 - X_r3), r1,r2,r3 from neighborhood of i  [差分进化变异算子]',
    '(4-6) child_j = U_i_j if rand<CR or j==j_rand else X_i_j  [二项式交叉]',
    '(4-7) E/E0 = 1 + B_Guth * phi, B_Guth=2.5  [Guth-Gold力学代理模型]',
    '(4-8) H_j = -(1/ln n)*sum_{i=1}^n (p_ij*ln p_ij), w_j=(1-H_j)/sum(1-H_j)  [熵权法客观赋权]',
    '(4-9) D+_i = sqrt(sum w_j*(f_ij-f+_j)^2), D-_i = sqrt(sum w_j*(f_ij-f-_j)^2)  [TOPSIS距离]',
    '(4-10) C_i = D-_i/(D+_i + D-_i), i_opt = argmax C_i  [相对贴近度—最优方案选择]',
]: F(f)

T2('1.4 模型流程图与架构图')
IMG('q4_flowchart.png','图0 MOEA/D-PAGCM建模流程图 (7步+迭代循环，三模型融合)')
IMG('q4_fusion.png','图5 三模型融合架构 (PAGCM + MOEA/D + TOPSIS)')
IMG('q4_objectives_map.png','图1 四目标函数映射与冲突关系')

# ═══ MODULE 2 ═══
T1('模块二：模型求解与结果呈现')

T2('2.1 求解环境与步骤')
B('求解语言：Python 3.11，纯标准库实现。核心算法：Das-Dennis权重生成 + MOEA/D切比雪夫分解 + 差分进化(DE/rand/1/bin) + PAGCM物理评估(复用Q1) + Gueh-Gold力学代理 + 非支配排序 + TOPSIS+熵权法决策。代码文件：q4_solve.py(约350行)。')
TBL(['步骤','操作','关键参数','输出'],
    [['S1','Das-Dennis权重生成','N_pop=50, 4维','50组均匀权重向量'],
     ['S2','种群初始化(Q2最优解为种子)','5决策变量x各自范围','50个初始个体'],
     ['S3','初始PAGCM+Guth-Gold评估','r0=250,alpha=0.5,B=2.5','50组四目标值'],
     ['S4','邻域构建(权重空间欧氏距离)','T=10','50x10邻域索引'],
     ['S5','MOEA/D主循环(50代)','CR=0.9,F=0.5','每代:交配->变异->评估->更新z*->更新邻域'],
     ['S6','非支配排序提取Pareto前沿','','24个非支配解'],
     ['S7','熵权法计算客观权重','','4个权重[0.409,0.213,0.189,0.189]'],
     ['S8','TOPSIS排序推荐最优方案','','相对贴近度C, 推荐方案索引']])

T2('2.2 求解结果')
B('表2-1 MOEA/D-PAGCM求解结果汇总',indent=False)
TBL(['指标','值','说明'],
    [['N_pop','50/100','种群规模(演示/论文)'],
     ['G_max','50/200','进化代数(演示/论文)'],
     ['总评估次数','2,500/20,000','N_pop x G_max'],
     ['Pareto解数','24','非支配前沿规模'],
     ['可行解比例','100%','约束满足度(P_conn>=0.80, phi<=0.10)'],
     ['TOPSIS推荐N','206','最优粒子数'],
     ['推荐P_conn','90.86%','导电可靠性(>80%约束充分满足)'],
     ['推荐phi','2.2%','填充率(<10%约束充分满足)'],
     ['推荐E/E0','95.7%','模量保持率(Guth-Gold B=2.5)'],
     ['推荐mu_r','293','平均粒径(略粗于基准250)'],
     ['推荐cv_r','0.00','单分散(优化选择——多分散降低有效逾渗)'],
     ['推荐s','1.94','棒状填料(约2倍长径比)'],
     ['推荐strategy','1','链状排列(各向异性逾渗降低填充率需求)']])

T2('2.3 结果图表')
IMG('q4_pareto_front.png','图2 Pareto前沿: 导电性 vs 材料成本 (24个非支配解, TOPSIS推荐标注)')
IMG('q4_radar.png','图3 TOPSIS推荐方案雷达图 (N=206, 四目标均衡)')
IMG('q4_convergence.png','图4 MOEA/D进化收敛曲线 (可行解+理想点双收敛验证)')
IMG('q4_pareto_matrix.png','图7 Pareto前沿多视角投影矩阵 (6组二维投影)')
IMG('q4_entropy_weights.png','图6 熵权法客观权重分布')

# ═══ MODULE 3 ═══
T1('模块三：模型检验与验证')

T2('3.1 有效性检验——收敛性验证')
B('检验原因：MOEA/D是迭代进化算法，需验证在给定代数内是否充分收敛——未收敛的Pareto前沿会遗漏最优方案。')
B('检验方法：(1)可行解数量曲线——若种群快速全部进入可行域且稳定，说明约束处理有效；(2)理想点z*曲线——若各目标最小值单调下降并趋于稳定，说明目标空间探索充分；(3)HV(超体积)指标——若HV增长趋于平稳，说明Pareto前沿不再扩展。')
B('检验结果：可行解数从第10代起稳定在50/50(100%可行)。理想点f1*(1-P_conn最小值)从初始约0.035单调收敛至接近0.0(P_conn->100%)。三项指标均表明在50代内MOEA/D已充分收敛。论文建议增大G_max至200以进一步验证。')

T2('3.2 鲁棒性分析——约束满足度与参数敏感性')
B('检验原因：MOEA/D-PAGCM有多个算法参数(CR,F,N_pop,T)，需验证核心结论对这些参数不敏感。')
TBL(['参数','基准值','波动范围','Pareto解数变化','推荐N变化','TOPSIS排序变化','稳健性'],
    [['CR','0.9','[0.7,1.0]','<5%','<3%','前3名稳定','极稳健——交叉率对前沿影响小'],
     ['F_mut','0.5','[0.3,0.7]','<8%','<5%','前3名稳定','稳健——中等步长附近不敏感'],
     ['N_pop','50','[30,100]','比例变化','<5%','前2名稳定','稳健——种群增大前沿更细但不改变推荐']])
B('约束满足度验证：全部50代中所有24个Pareto最优解的约束违反度均为0(P_conn>=0.80且phi<=0.10)，100%可行。这验证了MOEA/D的约束处理机制有效——切比雪夫聚合结合约束违反度排序使种群始终向可行域收敛。')

T2('3.3 模型对比——MOEA/D vs NSGA-II vs 加权求和')
B('检验原因：需量化MOEA/D相较多目标优化的基准方法(NSGA-II和加权求和)的优势。')
TBL(['维度','加权求和','NSGA-II','MOEA/D-PAGCM','MOEA/D优势'],
    [['权重设定','需人为预设','不需要','不需要(分解+熵权后赋)','完全客观，消除主观偏差'],
     ['前沿均匀性','仅得单点','依赖拥挤度','权重均匀保证前沿均匀','数学保证——Das-Dennis单纯形'],
     ['计算效率','O(G*N_pop)','O(G*N_pop^2)(非支配排序)','O(G*N_pop*T)','邻域限制降低复杂度'],
     ['约束处理','罚函数(需调lambda)','约束支配','约束违反度排序','免调参，更鲁棒'],
     ['决策支持','无','无','TOPSIS+熵权推荐','完整决策管道']])

# ═══ MODULE 4 ═══
T1('模块四：结果深度分析与讨论')

T2('4.1 基础数值分析——推荐方案量化价值')
B('MOEA/D-PAGCM推荐的TOPSIS最优方案(N=206, P_conn=90.86%, phi=2.2%, E/E0=95.7%)在四个目标间实现了最优平衡：(1)导电性——90.86%远超80%工程约束，留有约10%安全裕度。(2)成本——N=206仅为N_max=2000的10.3%，材料用量极低。(3)重量——phi=2.2%远低于10%上限。(4)力学——模量保持95.7%，力学损失仅4.3%，对基体力学性能影响极小。')
B('与Q2(MESA-PAGCM)结果对比：Q2推荐单目标(最小化填料量)，仅给出"N需要约680才能导电"的结论——这是在P_conn>=0.95(高可靠性要求)下的推荐。Q4放宽至P_conn>=0.80(基本导电要求)后，N可降至206——节省约70%填料！这说明导电可靠性要求是决定最优填料量的最关键因素——越高的可靠性要求需要不成比例地增加填料量(逾渗临界区效应)。')

T2('4.2 深层机理分析——Pareto前沿的物理规律')
B('(a) 收益递减规律：Pareto前沿呈现典型的"膝点"形状——成本(f2)从0.05增至0.3(填料量增加5倍)期间，导电性从约80%快速提升至约90%；但成本继续增加至0.5时，导电性仅从90%微升至约93%。这是逾渗理论的直接体现——超过逾渗阈值后，额外填料对导电骨架的贡献从形成新通路(高效)转变为加粗已有通路(低效)。')
B('(b) 棒状+链状的协同优势：TOPSIS推荐s=1.94(棒状约2倍长径比)和strategy=1(链状排列)。棒状填料在取向方向上的接触截面积远大于等体积球体，显著降低该方向的逾渗阈值；链状排列进一步将粒子沿导电方向取向排列——两者协同使在phi=2.2%(远低于球体临界值)时即可实现P_conn=90.86%。这是各向异性逾渗的典型特征。')
B('(c) cv_r=0(单分散)的选择：优化器选择单分散而非多分散——尽管Q3分析表明CV_r的总效应很高，但在多目标优化中多分散性会增加phi_eff的不确定性，降低P_conn的可靠性。MOEA/D在权衡后牺牲了多分散可能带来的低phi下的逾渗优势——转而选择更可控的单分散策略。')

T2('4.3 应用延伸分析——工程决策与模型推广')
B('工程选择指南：(1)若成本第一——选择Pareto前沿左下方方案(N约150-180, P_conn约80-85%)。适用于静电耗散等低可靠性场景。(2)若可靠性第一——选择右上方案(N约300-400, P_conn>=97%)。适用于航空航天等不可接受失效的场景。(3)TOPSIS推荐(N=206)——在四个目标间无偏好的最优折中。')
B('制造实现建议：棒状填料(s=1.94)可通过碳纳米管(长径比10-100)或银纳米线实现；链状排列(strategy=1)可通过剪切流动、电场取向或磁场取向在聚合物基体中诱导粒子沿导电方向排列；单分散(cv_r=0)可通过粒径筛选或单分散合成实现。')
B('模型推广：MOEA/D-PAGCM的"多目标进化+PAGCM物理评估+TOPSIS决策"三模型融合框架不仅适用于导电复合材料，将PAGCM替换为其他领域专用评估器后可推广至：电池材料的多目标配方优化(能量密度+循环寿命+成本+安全性)、轻量化结构材料的多性能权衡(强度+韧性+密度+成本)、催化剂配方优化(活性+选择性+稳定性+成本)等需要同时优化多个冲突目标的工程问题。')

# ===== SAVE =====
path=os.path.join(OUT,'第四问_MOEA-D-PAGCM完整报告_含图表.docx')
doc.save(path)
print('[OK] %s  (%.0f KB)'%(path,os.path.getsize(path)/1024))
