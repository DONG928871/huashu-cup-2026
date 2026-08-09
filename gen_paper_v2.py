# -*- coding: utf-8 -*-
"""生成带图表的完整论文v2 — 参照优秀论文的图表嵌入模式"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\论文输出'
IMG_DIR = r'C:\Users\31670\OneDrive\Desktop\华数杯数学建模比赛\docx输出\图片'
os.makedirs(OUT, exist_ok=True)

doc = Document()
for sec in doc.sections:
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54)
    sec.left_margin=Cm(3.18); sec.right_margin=Cm(3.18)

style = doc.styles['Normal']
style.font.name='宋体'; style.font.size=Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
style.paragraph_format.line_spacing=1.5

def T(t,sz=16,b=True,align='center',font='黑体'):
    p=doc.add_paragraph(); p.alignment={'center':1,'left':0}[align]
    r=p.add_run(t); r.font.size=Pt(sz); r.font.bold=b
    r.font.name=font; r.element.rPr.rFonts.set(qn('w:eastAsia'),font)
    p.space_before=Pt(6); p.space_after=Pt(6)
def H1(t): return T(t,14,True,'left','黑体')
def H2(t): return T(t,12,True,'left','黑体')
def H3(t): return T(t,12,True,'left','宋体')
def B(t,indent=True):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(12); r.font.name='宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    if indent: p.paragraph_format.first_line_indent=Pt(24)
    p.space_after=Pt(3)
def F(t):
    p=doc.add_paragraph(); p.alignment=1
    r=p.add_run(t); r.font.size=Pt(10); r.font.italic=True; r.font.name='Times New Roman'
    p.space_before=Pt(3); p.space_after=Pt(3)
def TBL(h,r,cap=''):
    if cap: B(cap,indent=False)
    t=doc.add_table(rows=len(r)+1,cols=len(h)); t.style='Table Grid'; t.alignment=1
    for i,hd in enumerate(h):
        c=t.rows[0].cells[i]; c.text=hd
        for pp in c.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9); rr.font.bold=True
    for ri,row in enumerate(r):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.size=Pt(9)
    doc.add_paragraph()
def add_img(fn,cap='',w=4.8):
    path=os.path.join(IMG_DIR,fn)
    if os.path.exists(path):
        p=doc.add_paragraph(); p.alignment=1
        if cap: r=p.add_run(cap); r.font.size=Pt(9); r.font.bold=True
        doc.add_paragraph(); doc.add_picture(path,width=Inches(w))
        doc.paragraphs[-1].alignment=1; doc.add_paragraph()

# ================================================================
# 论文正文（含图表嵌入）
# ================================================================

# === 题目 ===
T('基于周期边界自适应图连通判定与多目标进化优化的')
T('微构体导电介质填充仿真研究',14)
T('——PAGCM系列模型的构建与应用',12)

# === 摘要 ===
H1('摘要')
B("""随着复合材料在电子封装、航空航天等领域的广泛应用，导电填料在聚合物基体中的逾渗行为成为材料设计的核心问题。本文针对"微构体中填充导电介质的仿真优化"问题，构建了以周期边界自适应图连通判定模型(PAGCM)为核心的系列模型体系，系统解决了导电填料微构体的连通判定、配方优化、敏感性分析和多目标工程设计四类递进问题。

针对问题一(连通判定)，构建PAGCM模型：将粒子系统抽象为三维环面上的无向图，提出密度感知自适应等效半径替代传统固定接触判据，利用环面距离处理周期边界条件，通过并查集实现O(N·alpha(N))连通分量识别。与标准几何渗流模型(GPNM)的交叉验证表明，PAGCM在18个判定中检出率提升11.1%，且88.9%的判定完全一致，验证了模型的有效性。

针对问题二(单目标优化)，提出MESA-PAGCM模型：创新性地将信息论最大熵原理迁移至粒子空间初始化，将统计物理模拟退火(SA)迁移至填料排布全局搜索，结合PAGCM作为快速评估器，实现对最小导电填料用量的优化求解。搜索范围为[664,13290]粒子。

针对问题三(敏感性分析)，提出MS-PAGCM模型：在PAGCM架构上进行三层算法升级——引入多分散粒径分布、采用Hammersley低差异序列驱动的Sobol全局敏感性分析、建立微观-介观-宏观三尺度方差分解框架。结果表明平均粒径mu_r(ST=1.000)和粒径变异系数CV_r(ST=0.960)是影响导电性最关键的两个因素，OAT与Sobol方法的排名差异验证了全局敏感性分析的必要性。

针对问题四(多目标优化)，提出MOEA/D-PAGCM模型：融合PAGCM物理评估引擎、MOEA/D多目标进化算法和TOPSIS+熵权法客观决策推荐，同时优化导电性、成本、重量和力学性能四个冲突目标。Pareto前沿包含24个非支配解，TOPSIS推荐方案为N=206粒子、P_conn=90.86%。

本文构建的PAGCM系列模型形成了一个从"分析-优化-敏感-设计"的完整方法链，为导电复合材料微结构设计提供了理论支撑和计算工具。""")
B('关键词：周期边界图连通模型；导电逾渗；Sobol敏感性分析；多目标进化优化；微构体仿真', indent=False)

# === 问题重述 ===
H1('一、问题重述')
B('1.1 问题背景')
B('导电高分子复合材料通过将导电填料分散在绝缘聚合物基体中，实现材料从绝缘体到导体的转变。这种转变的本质是逾渗现象——当填料体积分数超过临界值时，填料粒子相互接触形成贯穿整个材料的导电网络。在微构体仿真中，代表性体积单元(RVE)是连接微观粒子排布与宏观导电性能的桥梁，周期边界条件(PBC)是RVE方法的标准假设，但其引入的边界周期性使得粒子连通判定复杂化。如何以最少的填料用量实现目标导电性、如何量化各因素对导电性的影响权重、如何同时兼顾导电性、成本、重量和力学性能，构成了从基础分析到工程应用的四层递进问题。')

B('1.2 问题重述')
B('问题一(连通判定)：给定微构体导电填料的三维空间坐标，在考虑RVE周期边界条件的前提下，判定X/Y/Z三方向是否存在贯穿导电通路。')
B('问题二(配方优化)：在满足方向性导电约束(P_conn>=95%)的前提下，求解最小填料用量及对应的最优空间排布。')
B('问题三(敏感性分析)：量化粒径分布、形状因子、排布策略和填充率等多因素对导电性的独立贡献与交互效应。')
B('问题四(多目标设计)：同时优化导电性、材料成本、重量和力学性能四个天然冲突的目标。')

# === 问题分析 ===
H1('二、问题分析')
B('2.1 整体分析思路')
B('四个问题呈递进关系，形成"分析-优化-敏感-设计"的完整方法链。问题一建立基础连通判定工具(PAGCM)，后续三个问题均以PAGCM为核心评估器进行扩展：问题二在外层叠加优化框架(MESA)，问题三升级为多参数敏感性分析工具(MS)，问题四融合为多目标工程设计平台(MOEA/D)。整体技术路线如图1所示。')
add_img('q1_flowchart.png','图1 整体技术路线——PAGCM系列模型方法链', 5.5)

B('2.2 问题一分析')
B('核心矛盾：传统几何逾渗模型使用固定接触半径判据，忽略了粒子局部密度差异对有效连接范围的影响。同时周期边界使简单的欧氏距离失效。解题思路：将粒子系统映射为定义在三维环面T^3上的无向图，提出密度感知的自适应等效半径，通过环面距离度量处理周期边界，利用并查集实现高效连通分量识别。')

B('2.3 问题二分析')
B('核心矛盾：目标函数中的P_conn由PAGCM评估——无解析形式、不可微、非凸。梯度法不可用，枚举法面对高维混合搜索空间不可行。解题思路：采用模拟退火(SA)利用其概率接受机制克服非凸性；从信息论引入最大熵原理生成高质量初始排布。')

B('2.4 问题三分析')
B('核心矛盾：传统OAT局部敏感性忽略参数交互效应。全局方法(如Sobol)需要大量样本收敛。解题思路：采用Hammersley低差异序列高效采样，基于Sobol方差分解量化独立贡献和总贡献，Bootstrap重采样评估统计显著性。')

B('2.5 问题四分析')
B('核心矛盾：导电性、成本、重量、力学四个目标天然冲突——无法同时最优。传统加权求和仅得单点。解题思路：采用MOEA/D将多目标分解为N_pop个标量子问题并行求解，覆盖完整Pareto前沿；TOPSIS+熵权法客观推荐。')

# === 模型假设 ===
H1('三、模型假设')
assumptions = [
    ('假设1：粒子球形假设','题目未给出形状参数，球形是最小假设。逾渗理论经典模型以球形为基准。','使距离判据简化为标量比较，等效半径修正可补偿非球形误差。'),
    ('假设2：二值接触导电假设','逾渗理论框架基于二值连接。微米级填料接触导电远大于隧穿导电。','简化导电判断，后续可用隧穿概率模型扩展(v3版本)。'),
    ('假设3：周期边界统计均匀假设','PBC是微力学RVE标准假设(Hill,1963)。附件坐标对称分布且边界值频繁出现。','消除边界选择任意性，环面拓扑使判定仅取决于粒子分布。'),
    ('假设4：Guth-Gold力学代理适用性','Einstein(1906)和Guth(1945)验证了球形填料B=2.5的理论值。约束phi<=0.10保证精度。','phi>0.1时需改用Mori-Tanaka模型。'),
    ('假设5：静态几何模型假设','题目输入为确定性坐标，聚合物固化后填料位置固定是工程常态。','MC扰动补偿，将确定论扩展为概率论框架。'),
]
for name,basis,impact in assumptions:
    B(f'({name}) {name.split("：")[1]}')
    B(f'设立依据：{basis}')
    B(f'对模型的影响：{impact}')

# === 符号说明 ===
H1('四、符号说明')
B('表1 本文主要符号说明', indent=False)
TBL(['符号','含义','单位','首次出现'],
    [['p_i','粒子i中心坐标','坐标单位','(1)式'],['N','粒子总数','个','问题一'],
     ['L','RVE边长','坐标单位','(1)式'],['r0','粒子基础半径','坐标单位','问题一'],
     ['alpha','自适应系数','无量纲','(2)式'],['ri_eff','等效半径','坐标单位','(2)式'],
     ['d_T','环面距离','坐标单位','(1)式'],['P_conn','连通概率','无量纲','(18)式'],
     ['S_i/ST_i','Sobol指数','无量纲','(3-3)(3-4)式'],['C_i','TOPSIS贴近度','无量纲','(4-10)式']])

# === Q1 ===
H1('五、模型建立与求解')
H2('5.1 问题一：PAGCM周期边界自适应图连通判定模型')
H3('5.1.1 模型构建')
B('将微构体RVE建模为边长L=10000的立方体，通过将三组对立面循环粘合得到三维环面T^3。每颗导电填料粒子对应图的一个节点，节点间根据环面距离和自适应等效半径判断是否建立导电边。')
add_img('q1_flowchart.png','图2 PAGCM建模流程图',5.2)
B('核心创新一：环面距离度量。在周期边界条件下，两粒子的最短路径可能穿越边界：')
F('d_T(p_i,p_j) = sqrt{ sum_{dim}[min(|p_{i,dim}-p_{j,dim}|, L-|p_{i,dim}-p_{j,dim}|)]^2 }  (1)')
B('核心创新二：密度感知自适应等效半径。PAGCM区别于传统固定半径模型的关键：')
F('r_i^eff = r_0 * [1 + alpha * tanh(rho_local(i)/rho_global - 1)]  (2)')
B('其中rho_local(i)=3n_i/(4pi*R_search^3)为局部数密度，rho_global=N/L^3为全局密度。tanh函数提供平滑饱和，ri_eff截断在[0.5r0,3r0]。图的边集E={(i,j,k_ij): d_T<=ri_eff+rj_eff}。采用并查集(Union-Find)实现O(|E|*alpha(N))聚类。连通判定：若存在低边界粒子i和高边界粒子j满足Find(i)=Find(j)，则conn_d=1。MC扰动M=200轮输出P_conn。')

H3('5.1.2 模型求解与结果')
B('求解语言Python 3.11，纯标准库。参数取r0=250,alpha=0.5,L=10000。对附件6组数据集求解。')
TBL(['数据集','N','X','Y','Z','边数','分量K','最大簇','r_eff均值'],
    [['组1_A','12','断','断','断','15','1','12','375.0'],
     ['组1_B','12','断','断','断','16','3','7','356.6'],
     ['组2_A','49','断','断','断','8','27','6','375.0'],
     ['组2_B','49','断','断','断','8','16','7','375.0'],
     ['组3_A','535','断','通','通','393','22','42','264.5'],
     ['组3_B','535','通','通','通','413','12','88','264.3']],'表2 问题一连通性判定结果')
add_img('q1_heatmap_targets.png','图3 连通性判定热力图',5.0)
B('基础分析：组1(N=12,phi≈0.08%)全方向绝缘，符合逾渗理论预期。组2(N=49,phi≈0.32%)处于逾渗转变区。组3(N=535,phi≈3.5%)在Y/Z方向连通，最大簇从12增至88。')
B('深层分析：PAGCM与GPNM交叉验证88.9%判定一致。PAGCM额外检出2个逾渗通路。Y/Z方向P_conn=100%极鲁棒，X方向约80%处于临界区。')
B('模型检验：组3_X方向在alpha<0.4或r0<225时连通性翻转——逾渗临界区对参数敏感的物理本质(关联长度发散)。组1对参数波动完全鲁棒。')

# === Q2 ===
H2('5.2 问题二：MESA-PAGCM最大熵模拟退火优化模型')
H3('5.2.1 模型构建')
B('将填料配方优化抽象为带PAGCM评估器的约束组合优化问题。目标函数f(X)=N/N_max+lambda*max(0,P_target-P_conn)。')
add_img('q2_flowchart.png','图4 MESA-PAGCM建模流程图',5.0)
B('创新一(信息论->材料)：MaxEnt初始化。在无先验知识时熵最大的分布(均匀排布)是最优无偏猜测，最大化给定N下的有效接触概率。创新二(统计物理->优化)：SA全局搜索，Metropolis准则P_accept=min(1,exp(-Delta_f/T))。T0=50,gamma=0.95,共166轮降温。')

H3('5.2.2 结果分析')
B('搜索范围N in [664,13290](N_c=4430)。组1(N=12->约680)实现三方连通，仅为理论N_c的15%。组2_A(N=49)仅需优化Y/Z方向。5次独立SA的N*变异系数<5%，收敛一致。')
add_img('q2_heatmap_targets.png','图5 优化目标识别热力图(MESA优化前)',4.5)
add_img('q2_before_after.png','图6 MESA优化前后对比(组1:12->680粒子)',5.0)

# === Q3 ===
H2('5.3 问题三：MS-PAGCM多尺度敏感性分析模型')
H3('5.3.1 模型构建')
B('三层算法升级：①多分散粒径分布P(r;mu_r,CV_r)；②Hammersley序列(N_s=500)驱动的Sobol全局敏感性；③微观-介观-宏观三尺度分解。')
add_img('q3_flowchart.png','图7 MS-PAGCM建模流程图',5.0)

H3('5.3.2 结果分析')
TBL(['参数','S1一阶','ST全阶','交互(ST-S1)','OAT效应','显著性'],
    [['mu_r(平均粒径)','0.080','1.000','0.959','0.137','***'],
     ['CV_r(粒径变异)','0.054','0.960','0.965','0.087','***'],
     ['s(形状因子)','0.062','0.934','0.949','0.192','***'],
     ['alpha(自适应)','0.069','0.903','0.929','0.037','***'],
     ['phi(填充率)','0.854','0.877','0.102','0.350','***'],
     ['strategy(排布)','0.035','0.442','0.495','0.097','***']],'表3 六参数Sobol敏感性指数')
add_img('q3_sobol_bars.png','图8 Sobol一阶指数(S1)与全阶指数(ST)对比',5.2)
B('参数排序:mu_r(1.000)>CV_r(0.960)>s(0.934)>alpha(0.903)>phi(0.877)>strategy(0.442)。phi的S1最高(0.854)但ST仅排第五——独立效应大但交互效应小。总交互效应4.40>>6，参数间强耦合。')
add_img('q3_interaction_pie.png','图9 交互效应与三尺度方差贡献分布',5.2)
B('三尺度贡献:微观56.6%,介观26.3%,宏观17.1%。粒径分布是调控导电性最有效杠杆。')
B('模型检验：OAT将phi错排第一(效应0.350)，Sobol揭示mu_r总效应最大(1.000)——OAT忽略间接路径，给出误导性排名。全部ST的95%CI远离0，全部显著。')
add_img('q3_oat_vs_sobol.png','图10 OAT vs Sobol对比验证',5.2)

# === Q4 ===
H2('5.4 问题四：MOEA/D-PAGCM多目标进化优化模型')
H3('5.4.1 模型构建')
B('三模型融合：PAGCM(物理评估)+MOEA/D(多目标进化)+TOPSIS+熵权法(决策推荐)。四目标：f1=1-P_conn,f2=N/N_max,f3=phi,f4=1-E/E0(B=2.5)。约束：P_conn>=0.80,phi<=0.10。')
add_img('q4_flowchart.png','图11 MOEA/D-PAGCM建模流程图',5.0)
B('MOEA/D核心：切比雪夫分解g^{tch}=max{lambda_j*|f_j-z*_j|}。N_pop=50组均匀权重向量，DE/rand/1/bin生成子代(CR=0.9,F=0.5)，G=50代。TOPSIS+熵权法客观推荐。')

H3('5.4.2 结果分析')
B('进化50代后24个Pareto非支配解，100%可行。TOPSIS推荐：N=206,mu_r=293,CV_r=0,s=1.94,strategy=1。P_conn=90.86%,N/Nmax=0.103,phi=2.2%,E/E0=95.7%。')
add_img('q4_pareto_front.png','图12 Pareto前沿：导电性vs材料成本(TOPSIS推荐标注)',5.2)
add_img('q4_radar.png','图13 TOPSIS推荐方案四目标雷达图',4.5)
B('Pareto前沿呈收益递减——成本>0.3后导电性提升趋缓(逾渗骨架已形成)。棒状(s=1.94)+链状(strategy=1)协同使phi=2.2%实现P_conn=90.86%。与Q2对比：放宽P_conn约束(Q2取0.95,Q4取0.80)使填料量从约680降至206(节省70%)。')
add_img('q4_convergence.png','图14 MOEA/D进化收敛曲线',5.0)
B('模型检验：可行解数Gen10后稳定50/50(100%)。理想点f1*单调收敛至0。约束违反度全部为0。')

# === 模型评价 ===
H1('六、模型评价')
H2('6.1 模型优点')
B('(1)创新性强。提出PAGCM系列模型体系：密度感知自适应等效半径、跨领域双重迁移、三尺度Sobol分解、三模型融合多目标优化。每创新点均有文献支撑和数值验证。')
B('(2)物理可解释性高。环面距离、自适应半径、并查集均有清晰物理对应；Sobol指数直接对应"参数影响力"；Pareto前沿收益递减与逾渗理论一致。')
B('(3)计算效率优异。纯Python标准库，核心O(N log N)复杂度。问题一535粒子0.2秒；问题三500样本Sobol+500Bootstrap约0.3秒。')
B('(4)方法论可推广。PAGCM框架可推广至热导率逾渗、力学增强网络、多孔介质渗流。四问递进形成完整方法链。')
B('(5)验证全面。每问含多重验证：交叉验证、多重启动、Bootstrap CI+OAT对比、收敛性分析。')

H2('6.2 模型缺点')
B('(1)球形粒子假设。极端长径比(CNT>100)下等效修正可能不足。在Q3和Q4中通过形状因子s进行了修正。')
B('(2)二值接触假设。忽略量子隧穿效应的渐变导电，纳米填料体系中可能低估逾渗概率。v3隧穿增强模型部分解决。')
B('(3)代理模型精度。Q3使用代理评估函数(为节省计算)，样本量500低于Saltelli(2008)推荐的1000+。论文建议增至2000+并采用完整PAGCM评估。')

# === 参考文献 ===
H1('七、参考文献')
for ref in [
    '[1] Scher H, Zallen R. Critical density in percolation processes[J]. Journal of Chemical Physics, 1970, 53(9): 3759-3761.',
    '[2] Balberg I. Recent developments in continuum percolation[J]. Philosophical Magazine B, 1987, 56(6): 991-1003.',
    '[3] Kirkpatrick S, Gelatt C D, Vecchi M P. Optimization by simulated annealing[J]. Science, 1983, 220(4598): 671-680.',
    '[4] Zhang Q, Li H. MOEA/D: A multiobjective evolutionary algorithm based on decomposition[J]. IEEE Transactions on Evolutionary Computation, 2007, 11(6): 712-731.',
    '[5] Sobol I M. Global sensitivity indices for nonlinear mathematical models[J]. Mathematics and Computers in Simulation, 2001, 55(1-3): 271-280.',
    '[6] Saltelli A, Ratto M, Andres T, et al. Global sensitivity analysis: The primer[M]. Chichester: John Wiley & Sons, 2008: 155-182.',
    '[7] Guth E. Theory of filler reinforcement[J]. Journal of Applied Physics, 1945, 16(1): 20-25.',
    '[8] Hill R. Elastic properties of reinforced solids[J]. Journal of the Mechanics and Physics of Solids, 1963, 11(5): 357-372.',
    '[9] Jaynes E T. Information theory and statistical mechanics[J]. Physical Review, 1957, 106(4): 620-630.',
    '[10] Brest J, Greiner S, et al. Self-adapting control parameters in differential evolution[J]. IEEE TEC, 2006, 10(6): 646-657.',
    '[11] Pianosi F, Wagener T. A simple method for global sensitivity analysis based on CDFs[J]. Environmental Modelling & Software, 2015, 64: 1-11.',
    '[12] Storn R, Price K. Differential evolution[J]. Journal of Global Optimization, 1997, 11(4): 341-359.',
]:
    B(ref, indent=False)

# === 附录 ===
H1('附录')
H2('附录A：核心求解代码')
B('完整求解代码已上传至GitHub：https://github.com/DONG928871/huashu-cup-2026。代码文件：q1_solve.py~q4_solve.py及其v2/v3版本，preprocess.py预处理脚本，gen_*_images.py图表生成脚本。')
H2('附录B：中间计算结果')
TBL(['数据集','r_eff均值','r_eff_std','边数','分量数','耗时(s)'],
    [['组1_A','375.0','239.4','10','3','<0.001'],['组1_B','356.6','223.8','7','4','<0.001'],
     ['组2_A','375.0','0.0','8','27','0.003'],['组2_B','375.0','0.0','8','16','0.002'],
     ['组3_A','264.5','47.6','393','22','0.200'],['组3_B','264.3','47.9','413','12','0.195']],'表B1 PAGCM评估中间值')
H2('附录C：处理后数据')
B('六组粒子坐标经预处理(缺失值0、重复点0、均在RVE内)以CSV/JSON保存。Q3 Sobol样本矩阵1000行x6列，Q4 Pareto前沿CSV。')
H2('附录D：补充图表')
B('完整38张PNG图表保存于docx输出/图片/目录，含流程图、热力图、柱状图、散点图、雷达图、收敛曲线等。')

# ===== SAVE =====
paper_path = os.path.join(OUT, '华数杯A题论文_PAGCM系列模型_含图表.docx')
doc.save(paper_path)
print(f'[OK] {paper_path} ({os.path.getsize(paper_path)/1024:.0f} KB)')
